from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

doc = Document()

# ── 页面边距 ──
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2.5)

# ── 工具函数 ──
def set_font(run, size=11, bold=False, color=None, name_cn='宋体', name_en='Calibri'):
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    run.font.name = name_en
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name_cn)

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.name = 'Calibri'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        if level == 1:
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        elif level == 2:
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
        elif level == 3:
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
    return p

def add_para(doc, text, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    set_font(run)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    run = p.add_run(text)
    set_font(run)
    return p

def add_table_row(table, cells_data, is_header=False):
    row = table.add_row()
    for i, text in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(text)
        set_font(run, size=9.5, bold=is_header,
                 name_cn='宋体' if not is_header else '黑体')
        if is_header:
            shading = OxmlElement('w:shd')
            shading.set(qn('w:val'), 'clear')
            shading.set(qn('w:color'), 'auto')
            shading.set(qn('w:fill'), '2E74B5')
            cell._tc.get_or_add_tcPr().append(shading)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

# ══════════════════════════════════════════
# 封面
# ══════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title_p.add_run('イナズマイレブン クロス\n新作调研报告')
r.font.size = Pt(22)
r.font.bold = True
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
r.font.name = 'Calibri'
r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

doc.add_paragraph()
meta_p = doc.add_paragraph()
meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for line in ['调查时间：2026年6月',
             '主要数据源：game8.jp / appmedia.jp / Google Play',
             '市场参照：放置系手游市场趋势（以 AFK: Journey 为代表）']:
    r = meta_p.add_run(line + '\n')
    set_font(r, size=11, color=(0x60, 0x60, 0x60))

doc.add_page_break()

# ══════════════════════════════════════════
# 一、专有名词
# ══════════════════════════════════════════
add_heading(doc, '一、专有名词与对应机能', 1)
add_para(doc, '以下术语为游戏内核心专有名词，括号内为读法，最右列为 AFK: Journey 的对应概念。')
doc.add_paragraph()

terms = [
    ('TP（テクニカルポイント）', 'Technical Point',
     '对战中使用必殺技消耗的资源；命令战胜利积累，失败减少；每个技能有固定TP消耗值', 'Mana / Energy'),
    ('必殺技（ひっさつわざ）', '必杀技',
     '4种类型：シュート（射门）/ ドリブル（运球）/ ブロック（拦截）/ キーパー（扑救）；附带属性、TP、射程、暴击率参数', '技能 / Ultimate'),
    ('属性（4すくみ）', '风·林·火·山',
     '四属性循环克制：风克山→山克火→火克林→林克风；克属性时必殺技威力有加成', '元素 / 装甲克制'),
    ('ポジション', 'FW/MF/DF/GK',
     '前锋/中场/后卫/门将，决定场上职责与推荐技能类型', '职业 / Role'),
    ('覚醒（かくせい）', '觉醒',
     'ADVANCE → ADVANCE+ → TOP → TOP+ → LEGENDARY；提升基础数值并解锁/强化被动技能；需消耗覚醒魂和属性碎片', '超越 / Ascension'),
    ('覚醒魂（かくせいたましい）', '觉醒魂',
     '觉醒材料；来源：抽卡重复角色、活动商店、人脈ボード', '英雄碎片 / Soulstone'),
    ('人脈ボード', '人脉版',
     '网格式进度解锁系统，完成后可获得角色、道具；由对战与信頼度系统推进', '英雄感情板 / Bond Board'),
    ('信頼度（しんらいど）', '信任度',
     '出阵后自动积累；达里程碑奖励ダイヤボール（高级货币）；解锁人脈ボード节点', '好感度 / Friendship'),
    ('クロスシミュレーター', '体能训练器',
     '挂机自动收益系统；支持ハイライトモード倍速；产出角色培育材料（サプリ）', 'AFK报酬 / Idle Chest'),
    ('サプリ', '补给品',
     '达到关键等级（每10级）时消耗的特殊材料，解锁被动技能成长；仅来自クロスシミュレーター', '进阶石材'),
    ('秘伝書（ひでんしょ）', '秘传书',
     '将技能"教给"不具备该技能的角色；リミテッドバトル等用于填补战术空缺', '技能书 / Skill Tome'),
    ('ダイヤボール', '钻石球', '高级抽卡货币；由信頼度奖励、任务、活动获取', '钻石 / Diamond'),
    ('熱血ポイント', '热血值', 'PvP（対戦）奖励货币，随段位提升获取量增加', '竞技币 / Arena Coin'),
    ('コアメンバー', '核心成员', '队伍中5名主要角色，等级最低者决定全队上限；强制均衡培养设计', '共鸣等级限制'),
    ('シュートチェイン', '连射', '部分射门技能特性；允许多角色连续射门叠加；技能标注「可/不可」', '连击 / Chain'),
    ('対戦', 'PvP',
     '排名赛制：NORMAL→GROWING→ADVANCED→TOP→LEGENDARY→HERO；每日免费5次；奖励按日/周/赛季三级结算', '竞技场 / Arena'),
]

table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
table.autofit = False
widths = [Cm(4.5), Cm(2.5), Cm(8.0), Cm(3.5)]
for i, w in enumerate(widths):
    for cell in table.columns[i].cells:
        cell.width = w

add_table_row(table, ['专有名词', '读法 / 简称', '机能说明', 'AFK: Journey 对应'], is_header=True)
for row_data in terms:
    add_table_row(table, list(row_data))

doc.add_paragraph()

# ══════════════════════════════════════════
# 二、活动玩法
# ══════════════════════════════════════════
add_heading(doc, '二、活动玩法规则（3个案例）', 1)

# 活动1
add_heading(doc, '活动① サービス開始記念イベント（服务开始纪念活动）', 2)
add_para(doc, '期间：2026/6/9 ～ 7/9　|　来源：game8.jp/inazuma-cross/788218', indent=True)
add_heading(doc, '核心机制', 3)
add_bullet(doc, '进行活动专用关卡获取「活动代币」')
add_bullet(doc, '代币在限时商店兑换奖励（优先级：タペストリー 2000枚、秘伝書 1500枚）')
add_bullet(doc, '配套「個人人脈ボード」，完成全节点可确定获得角色「栗松鉄平」并全觉醒，无需抽卡')
add_heading(doc, '设计特点', 3)
add_para(doc, '确定奖励角色通过"版面完成"而非概率抽卡获取，降低零氪门槛，但需持续刷关推进。结构与 AFK: Journey"限时剧情活动+奖励兑换"高度类似，差异在于角色获取方式更具确定性。', indent=True)

doc.add_paragraph()

# 活动2
add_heading(doc, '活动② リミテッドサッカーバトル（限定足球对战）', 2)
add_para(doc, '来源：game8.jp/inazuma-cross/789626', indent=True)
add_heading(doc, '核心机制', 3)
add_bullet(doc, '编队受限：只能使用特定派系角色（雷门、帝国、日本代表等），同派系至少需5名角色才可挑战')
add_bullet(doc, '阵型限制：可用阵型仅剩2种（普通模式无此限制）')
add_bullet(doc, '中央通道适性成为关键评估指标；边路型角色使用受限')
add_bullet(doc, 'GK补救：派系内无合适门将时，可用「秘伝書」跨职业学习守门技能（资源消耗较高）')
add_bullet(doc, '解锁条件：通关故事第3章第5战后')
add_heading(doc, '奖励', 3)
add_para(doc, '通关解锁新监督角色（提供新阵型）、监督经验道具、ガチャチケット', indent=True)
add_heading(doc, '设计特点', 3)
add_para(doc, '强制限定编队迫使玩家横向培养角色而非单押一队，与 AFK: Journey"英雄装备特殊关卡"思路相近，但派系划分呼应了原作 IP 的队伍对抗感。', indent=True)

doc.add_paragraph()

# 活动3
add_heading(doc, '活动③ 日本代表2026プロジェクト（国家队2026企划）', 2)
add_para(doc, '期间：2026/6/14 ～ 6/30　|　来源：game8.jp/inazuma-cross/790268', indent=True)
add_heading(doc, '核心机制', 3)
add_bullet(doc, '限时限定抽卡：「[日本代表2026]円堂守」「[日本代表2026]鬼道有人」两名限定角色')
add_bullet(doc, '多种制服造型可分别收集（主场/客场/GK服）')
add_bullet(doc, '大厅装饰物：登录期间可获取纪念装饰品用于「クロスラウンジ」个人空间')
add_bullet(doc, '活动游戏量轻：以登录 + 限定抽卡为主，无重度刷关')
add_heading(doc, '设计特点', 3)
add_para(doc, 'IP联动型活动，强调收藏性与情怀触发，实际游戏深度较低。核心消费驱动力为角色造型而非战力，与 AFK 限定皮肤活动形式相似。', indent=True)

doc.add_paragraph()

# ══════════════════════════════════════════
# 三、技能数据
# ══════════════════════════════════════════
add_heading(doc, '三、技能数据摘要', 1)
add_para(doc, '数据来源：Kinoko Corpus 数据库（game8.jp 已收录 44 条技能，为目前已知全量或接近全量）')

doc.add_paragraph()
add_heading(doc, '技能结构', 2)
add_para(doc, '所有技能均采用统一格式，字段极简：')
code_p = doc.add_paragraph()
code_p.paragraph_format.left_indent = Cm(1)
r = code_p.add_run('【イナイレクロス：技名】 種別：X 属性：X TP：X クリ：X% 効果：射程N + 威力N的X技')
r.font.name = 'Courier New'
r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
r.font.size = Pt(9.5)

add_para(doc, '注：所有"效果"字段仅为数值参数，无任何游戏机制说明文字。')
doc.add_paragraph()

add_heading(doc, '数据分布概览', 2)
t2 = doc.add_table(rows=1, cols=4)
t2.style = 'Table Grid'
add_table_row(t2, ['种别', '数量', 'TP范围', '暴击率范围'], is_header=True)
skill_rows = [
    ('シュート（射门）', '20条', '30–60', '4%–11%'),
    ('ドリブル（运球）', '13条', '30–60', '5%–26%'),
    ('ブロック（拦截）', '8条',  '30–50', '9%–19%'),
    ('キーパー（扑救）', '3条',  '40–60', '10%–25%'),
]
for row_data in skill_rows:
    add_table_row(t2, list(row_data))

doc.add_paragraph()
add_heading(doc, '最强技（按威力）', 2)
add_bullet(doc, '射门：爆熱スクリュー（威力199，TP60，火属性）')
add_bullet(doc, '运球：キラーフィールズ（威力158，TP60，林属性）')
add_bullet(doc, '拦截：ザ・マウンテン（威力172，TP50，山属性）')
add_bullet(doc, '扑救：ムゲン・ザ・ハンド（威力177，TP60，林属性）')

doc.add_paragraph()
add_heading(doc, '技能设计特征', 2)
add_para(doc, '与 AFK: Journey 等 RPG 相比，イナイレクロス技能设计极度数值化：无状态效果、无持续技、无条件触发文字，完全由「威力×TP×属性×射程×暴击率」5个参数定义。这与足球运动本身的直观性高度吻合，牺牲了策略深度换取了规则透明度。')

doc.add_paragraph()

# ══════════════════════════════════════════
# 四、玩家评价（NLP情感分析）
# ══════════════════════════════════════════
add_heading(doc, '四、玩家评价（NLP情感分析）', 1)

note_p = doc.add_paragraph()
note_p.paragraph_format.left_indent = Cm(0.5)
shading = OxmlElement('w:shd')
shading.set(qn('w:val'), 'clear')
shading.set(qn('w:color'), 'auto')
shading.set(qn('w:fill'), 'E8F4FD')
note_p._p.get_or_add_pPr().append(shading)
r = note_p.add_run('数据说明：Twitter/X（HTTP 402）、App Store（API封锁）、5ch（HTTP 403）均无法抓取。本节基于 Google Play 日语最新评价 478 条（采样方式：NEWEST，按发布时间倒序；2026年6月抓取）。Google Play 总评价数为 7,372 件，公开 API 每次最多返回约 478 条，本样本约占总量 6.5%，为非随机抽样，以下结论为趋势性参考，不代表全量分布。分析工具：asari（日语NLP情感分析库）。')
set_font(r, size=9.5, color=(0x1F, 0x49, 0x7D))

doc.add_paragraph()

add_heading(doc, '多平台评分概览', 2)
t_platform = doc.add_table(rows=1, cols=4)
t_platform.style = 'Table Grid'
add_table_row(t_platform, ['平台', '评分总数', '平均分', '文字评论NLP分析'], is_header=True)
platform_rows = [
    ('Google Play', '7,372 件', '4.5 / 5.0', '478条（NEWEST样本，约占6.5%）'),
    ('App Store（日本）', '37,754 件', '4.74 / 5.0', '无法获取（Apple API限制）'),
]
for row_data in platform_rows:
    add_table_row(t_platform, list(row_data))

add_para(doc, 'App Store 评分数量是 Google Play 的 5 倍，且平均分更高，说明 iOS 用户满意度整体优于 Android 用户，或 iOS 用户对该 IP 的情感认同更强。文字评论仅能从 Google Play 获取，以下 NLP 分析基于此数据。', indent=True)

doc.add_paragraph()

add_heading(doc, 'Google Play 详细数据', 2)
t_nlp = doc.add_table(rows=1, cols=2)
t_nlp.style = 'Table Grid'
add_table_row(t_nlp, ['指标', '数值'], is_header=True)
nlp_rows = [
    ('分析样本数', '478 件（NEWEST排序，非随机）'),
    ('平均星评', '3.66 / 5.0'),
    ('NLP正面情感', '330件（样本中约69%，趋势参考）'),
    ('NLP负面情感', '148件（样本中约31%，趋势参考）'),
]
for row_data in nlp_rows:
    add_table_row(t_nlp, list(row_data))

doc.add_paragraph()

add_heading(doc, '星评分布', 2)
t_star = doc.add_table(rows=1, cols=3)
t_star.style = 'Table Grid'
add_table_row(t_star, ['星评', '件数', '占比'], is_header=True)
star_rows = [
    ('★5', '232', '48.5%'),
    ('★4', '64',  '13.4%'),
    ('★3', '52',  '10.9%'),
    ('★2', '46',  '9.6%'),
    ('★1', '84',  '17.6%'),
]
for row_data in star_rows:
    add_table_row(t_star, list(row_data))
add_para(doc, '注：两极化明显——★5接近半数，★1占17.6%，IP粉丝高度满意，非粉玩家强烈不满。', indent=True)

doc.add_paragraph()

add_heading(doc, '正面情感高频主题', 2)
t_pos = doc.add_table(rows=1, cols=3)
t_pos.style = 'Table Grid'
add_table_row(t_pos, ['关键词', '频次', '解读'], is_header=True)
pos_rows = [
    ('面白（有趣）', '51', '最核心正评词，游戏性获认可'),
    ('ストーリー（剧情）', '41', '剧情评价两极（也出现在负面词中）'),
    ('最高', '36', '情感强烈的极端正评'),
    ('イナズマイレブン', '33', 'IP情怀是核心正评来源'),
    ('ガチャ（抽卡）', '32', '正面语境下抽卡体验尚可'),
    ('必殺技', '10', '技能演出受欢迎'),
]
for row_data in pos_rows:
    add_table_row(t_pos, list(row_data))

doc.add_paragraph()
add_heading(doc, '正面代表评价', 3)
add_bullet(doc, '★3：「楽しいです。ストーリーのキャラクターボイスがあると助かります。」')
add_bullet(doc, '★5：「面白い、イナズマイレブン最高」')

doc.add_paragraph()

add_heading(doc, '负面情感高频主题', 2)
t_neg = doc.add_table(rows=1, cols=3)
t_neg.style = 'Table Grid'
add_table_row(t_neg, ['关键词', '频次', '解读'], is_header=True)
neg_rows = [
    ('試合（比赛）', '30', 'AI操控比赛体验差'),
    ('ガチャ（抽卡）', '21', '抽卡概率/机制投诉多'),
    ('操作', '21', '不可手动操作是主要槽点'),
    ('オート（自动）', '15', '全自动让重度玩家失落感'),
    ('シュート（射门）', '15', 'AI射门决策差'),
    ('意味（无意义）', '9', '游戏内容感觉空洞'),
]
for row_data in neg_rows:
    add_table_row(t_neg, list(row_data))

doc.add_paragraph()
add_heading(doc, '负面代表评价', 3)
add_bullet(doc, '★1：「ガチャ出ず、試合長い、スキップ出来ず、ガチャにいたっては近年まれにみるくらいの確率詐欺。」')
add_bullet(doc, '★2：「AIは本当にバカ。どう考えても目の前の豪炎寺にパスすればいいのに頑なに染岡に蹴る。ほんまレベルファイブは名作をゴミにする才能がすごい！！！」')
add_bullet(doc, '★1：「相手の強さはインフレしまくって勝つことが出来なくなる。強いシュート持ってるのに撃たず相手TP消費できず。」')

doc.add_paragraph()

add_heading(doc, '关键洞察', 2)
add_bullet(doc, 'AI战斗系统是最大槽点：多条高置信度负评集中在"AI决策太差"——不在最优时机射门、无意义回传、主动浪费TP，与"足球游戏"的手感期待严重背离。')
add_bullet(doc, 'ガチャ信任度偏低：正面评价中ガチャ出现频率也高，多条负评指出"天井低但根本出不了"的概率感知问题，存在"天花板体验欺骗"风险。')
add_bullet(doc, '剧情评价两极分化：ストーリー同时出现在正负频词中——IP粉丝看到经典角色满意，但剧情质量（新角色塞入、台词风格偏差）是主要负评来源。')
add_bullet(doc, '★5中的隐藏不满：部分★5评价经NLP检出负面情感，属"情怀溢价"现象——星评高但对特定设计有明显不满。')

doc.add_paragraph()

# ══════════════════════════════════════════
# 五、文字设计亮点
# ══════════════════════════════════════════
add_heading(doc, '五、文字设计亮点', 1)

add_heading(doc, '1. 属性名选择回避了抽象符号，强化世界观代入感', 2)
add_para(doc, '游戏使用「風・林・火・山」而非常见的「水・火・风・土」或色彩标签命名属性体系。「風林火山」来自武田信玄的军旗战术——"疾如风，徐如林，侵掠如火，不动如山"——这四个字在日本文化语境中有极强的武道与团队精神联想，与足球竞技气质相当契合。属性本身带叙事色彩，比纯色彩标签更有记忆点。')

add_heading(doc, '2. 「覚醒」段位使用全英文状态词，制造成就感层次', 2)
add_para(doc, 'ADVANCE → ADVANCE+ → TOP → TOP+ → LEGENDARY 的段位命名没有采用日文汉字数字（初級/中級/上級），而使用英文"升格词"。这种设计在视觉上让每次突破感更强烈——"LEGENDARY"作为终点词自带稀有感，比「最高級」更能激发玩家追求动机。')

add_heading(doc, '3. 「人脈ボード」命名精准呼应 IP 核心主题', 2)
add_para(doc, '大多数同类游戏把类似系统叫「絆ボード」或「育成マップ」，而本作选用「人脈（人际关系/人脉网络）」——这一词汇本身就是イナイレ系列的精神内核之一（"足球连接人与人"）。命名上的用心让系统功能与 IP 世界观直接挂钩，而不仅仅是一个数值解锁格子。')

add_heading(doc, '4. 「熱血ポイント」比「対戦コイン」情感浓度更高', 2)
add_para(doc, 'PvP货币命名为「熱血ポイント」而非通常的「対戦メダル」或「勝利コイン」。「熱血（热血）」是イナイレ系列最标志性的精神词汇，用它命名核心游戏资源，让每一场 PvP 对战都带上了情感意义，而不只是资源收割动作。这是将 IP 文化精髓渗透到 UI 层级的典型范例。')

doc.add_paragraph()

# ══════════════════════════════════════════
# 六、与AFK: Journeyの相似度重评
# ══════════════════════════════════════════
add_heading(doc, '六、与 AFK: Journey 的市场定位对比', 1)

note_p2 = doc.add_paragraph()
note_p2.paragraph_format.left_indent = Cm(0.5)
shading2 = OxmlElement('w:shd')
shading2.set(qn('w:val'), 'clear')
shading2.set(qn('w:color'), 'auto')
shading2.set(qn('w:fill'), 'E8F4FD')
note_p2._p.get_or_add_pPr().append(shading2)
r = note_p2.add_run('说明：AFK: Journey 作为当前放置系手游的市场标杆，用于衡量 イナイレクロス 在行业框架下的定位，而非直接竞品比较。')
set_font(r, size=9.5, color=(0x1F, 0x49, 0x7D))

doc.add_paragraph()
add_heading(doc, '相似度逐维度拆解', 2)

t_cmp = doc.add_table(rows=1, cols=4)
t_cmp.style = 'Table Grid'
add_table_row(t_cmp, ['维度', 'イナイレクロス', 'AFK: Journey', '相似度'], is_header=True)
cmp_rows = [
    ('战斗核心', '关键时刻手动选择释放哪个必殺技（消耗对应TP值），其余过程自动进行', '全自动，战中无任何操作', '❌ 完全不同'),
    ('挂机收益', 'クロスシミュレーター自动产出材料', 'Idle Chest自动产出', '✅ 相似'),
    ('角色成长路径', '覚醒段位（ADVANCE→LEGENDARY）＋碎片消耗', 'Ascension＋Soulstone', '✅ 相似'),
    ('关系/羁绊板', '人脈ボード网格解锁', 'Bond Board', '✅ 相似'),
    ('抽卡＋编队', '有', '有', '✅ 相似'),
    ('竖屏卡面UI风格', '有', '有', '✅ 相似'),
    ('战斗中阵型策略', '有（フォーメーション影响站位）', '有（Formation）', '✅ 相似'),
]
for row_data in cmp_rows:
    add_table_row(t_cmp, list(row_data))

doc.add_paragraph()
add_heading(doc, '定位总结', 2)
add_para(doc, '元游戏层与放置系市场高度接轨——挂机收益、抽卡、成长体系、竖屏UI框架均符合以 AFK: Journey 为代表的放置系手游行业标配，说明 イナイレクロス 在商业模式和用户习惯引导上有意向该市场靠拢，降低了新用户的认知门槛。')
doc.add_paragraph()
add_para(doc, '核心玩法保留了 IP 本身的差异化——战斗大部分自动进行，但在射门/运球/拦截等关键时刻会暂停，由玩家选择释放哪个必殺技（或选择传球保留TP）。决策深度有限，更接近"关键节点的辅助介入"而非完整指令制，但相比 AFK: Journey 的纯自动播放，确实多了一层互动参与感。')
doc.add_paragraph()
add_heading(doc, '市场定位结论', 3)
add_para(doc, 'イナイレクロス は一款以放置系手游为外壳、以指令制足球对战为内核的 IP 向手游。外壳降低入门门槛，内核服务 IP 粉丝的情感期待，两者服务于不同层次的玩家需求。', indent=True)

# ══════════════════════════════════════════
# 七、放置コンテンツ深度評価
# ══════════════════════════════════════════
add_heading(doc, '七、放置コンテンツ深度評価（クロスシミュレーター詳細）', 1)

note_p3 = doc.add_paragraph()
note_p3.paragraph_format.left_indent = Cm(0.5)
shading3 = OxmlElement('w:shd')
shading3.set(qn('w:val'), 'clear')
shading3.set(qn('w:color'), 'auto')
shading3.set(qn('w:fill'), 'E8F4FD')
note_p3._p.get_or_add_pPr().append(shading3)
r = note_p3.add_run('调查依据：游戏内ヘルプ画面（截图实测，2026年6月）')
set_font(r, size=9.5, color=(0x1F, 0x49, 0x7D))

doc.add_paragraph()
add_heading(doc, '系统概要', 2)
add_para(doc, 'クロスシミュレーター是本作的放置收益入口。「おまかせモード」开启后自动攻略シミュレーションステージ，产出角色培育素材（サプリ等）。')
doc.add_paragraph()
add_heading(doc, '关键限制', 3)
add_bullet(doc, '报酬积累上限：24小时（超过部分消失）')
add_bullet(doc, 'アプリを閉じると即停止——不是真正的后台放置，需保持应用运行')

doc.add_paragraph()
img_p1 = doc.add_paragraph()
img_p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
img_run1 = img_p1.add_run()
img_run1.add_picture(
    r'd:\游戏脚本\Kinoko_Corpus\screenshots\sim.jpg',
    width=Cm(7)
)
cap1 = doc.add_paragraph('图④ クロスシミュレーター ヘルプ画面（实机截图）。红框①：24時間上限；红框②：关闭App即停止。')
cap1.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(cap1.runs[0] if cap1.runs else cap1.add_run(''), size=9, color=(0x60, 0x60, 0x60))

doc.add_paragraph()
add_heading(doc, '加速系统', 2)

t_acc = doc.add_table(rows=1, cols=2)
t_acc.style = 'Table Grid'
add_table_row(t_acc, ['项目', '规格'], is_header=True)
acc_rows = [
    ('加速1次效果', '相当于2小时的成果受け取り'),
    ('无料补充', '1日1回'),
    ('追加加速（付费）', '1日最大5回（ダイヤボール消费）'),
    ('振替加速', '未使用分最多积攒7日分，可任意时机使用'),
]
for row_data in acc_rows:
    add_table_row(t_acc, list(row_data))

doc.add_paragraph()
img_p2 = doc.add_paragraph()
img_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
img_run2 = img_p2.add_run()
img_run2.add_picture(
    r'd:\游戏脚本\Kinoko_Corpus\screenshots\omakase.jpg',
    width=Cm(7)
)
cap2 = doc.add_paragraph('图⑤ 加速・振替加速 ヘルプ画面（实机截图）。橙框：加速1回=2時間・1日1回无料；绿框：振替加速7日繰り越し（AFK: Journey类似设计）。')
cap2.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(cap2.runs[0] if cap2.runs else cap2.add_run(''), size=9, color=(0x60, 0x60, 0x60))

doc.add_paragraph()
add_heading(doc, '振替加速补充规则', 3)
add_bullet(doc, '超过7日分的加速在日期更新时自动废弃')
add_bullet(doc, '振替加速获得的报酬以「未使用加速当日的最终关卡进度」为基准，不随后续关卡推进更新')

doc.add_paragraph()
add_heading(doc, '与 AFK: Journey 对比', 2)

t_afk = doc.add_table(rows=1, cols=3)
t_afk.style = 'Table Grid'
add_table_row(t_afk, ['比较轴', 'AFK: Journey', 'イナイレクロス'], is_header=True)
afk_rows = [
    ('后台放置（关闭App）', '✅ バックグラウンドでも継続', '❌ アプリを閉じると即停止'),
    ('报酬积累上限', '数日分（複数日）', '24時間（超過分は消滅）'),
    ('加速繰り越し', '✅ 未使用加速の繰り越し有り', '✅ 振替加速（最大7日分）'),
    ('放置的地位', 'コアループ（メイン進行）', 'サブコンテンツ（育成補助）'),
]
for row_data in afk_rows:
    add_table_row(t_afk, list(row_data))

doc.add_paragraph()
add_heading(doc, '评估结论', 2)
add_bullet(doc, '「振替加速（7日繰り越し）」是与 AFK: Journey 放置设计语言最相近的具体设计点。')
add_bullet(doc, '但放置深度总体有限：24小时上限偏短，关闭App即停止，不具备真正挂机体验，放置产出仅为辅助育成素材。')
add_bullet(doc, '综合评估：イナイレクロス 的放置设计属于「IP游戏附加的放置入门层」，并非独立深度的放置玩法；优先方向为收集放置相关用词（クロスシミュレーター、振替加速、おまかせ、加速 等）作为语料参照。')

doc.add_paragraph()
doc.add_paragraph()
end_p = doc.add_paragraph()
end_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = end_p.add_run('— 报告完成 —')
set_font(r, size=10, color=(0x80, 0x80, 0x80))

out_path = r'd:\游戏脚本\Kinoko_Corpus\イナイレクロス_竞品调研报告.docx'
doc.save(out_path)
print(f'Saved: {out_path}')
